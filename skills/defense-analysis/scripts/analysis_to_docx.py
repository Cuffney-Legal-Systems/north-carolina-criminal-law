#!/usr/bin/env python3
"""
analysis_to_docx.py — render a defense-analysis working draft (markdown) into
the delivered .docx, preserving the conventions in
reference/analysis-template.md.

    python3 analysis_to_docx.py <input.md> <output.docx>

Handles: fenced blocks (header block / closing notice), `>` best-issue flags,
##/###/#### headings, pipe tables (real Word tables with a repeating header
row — five columns for element tables, three for the bottom line), `-` bullets,
`1.` numbered action items, and inline **bold** / *italic* / `code`.

Every page carries the ATTORNEY WORK PRODUCT legend in the footer.

Requires python-docx.
"""
import re
import sys

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor

BODY, MONO = "Times New Roman", "Consolas"
LEGEND = "ATTORNEY WORK PRODUCT — PRIVILEGED"

# Column widths by column count; each row sums to 6.5" (letter less 1" margins).
WIDTHS = {
    5: [0.32, 1.82, 1.70, 0.85, 1.81],   # #, Element, State's proof, Strength, Defense attack
    4: [0.38, 2.32, 0.95, 2.85],         # legacy 4-column element table
    3: [2.20, 1.75, 2.55],               # Charge, State's proof, Defense posture
    2: [1.80, 4.70],
}

INLINE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
ORDERED = re.compile(r"^(\d+)\.\s+")


# ---------- low-level XML helpers ----------
def _sub(parent, tag, **attrs):
    el = OxmlElement(tag)
    for k, v in attrs.items():
        el.set(qn(k), v)
    parent.append(el)
    return el


def shade(el, fill):
    pr = el.get_or_add_tcPr() if el.tag.endswith("}tc") else el.get_or_add_pPr()
    _sub(pr, "w:shd", **{"w:val": "clear", "w:color": "auto", "w:fill": fill})


def border(par, edges, size="6", color="888888", space="6"):
    pr = par._p.get_or_add_pPr()
    bd = pr.find(qn("w:pBdr"))
    if bd is None:
        bd = _sub(pr, "w:pBdr")
    for e in edges:
        _sub(bd, f"w:{e}", **{"w:val": "single", "w:sz": size,
                              "w:space": space, "w:color": color})


def repeat_header(row):
    _sub(row._tr.get_or_add_trPr(), "w:tblHeader", **{"w:val": "true"})


def page_field(par):
    for instr in ("PAGE", "NUMPAGES"):
        if instr == "NUMPAGES":
            par.add_run(" of ")
        r = par.add_run()._r
        _sub(r, "w:fldChar", **{"w:fldCharType": "begin"})
        t = _sub(r, "w:instrText", **{"xml:space": "preserve"})
        t.text = f" {instr} "
        _sub(r, "w:fldChar", **{"w:fldCharType": "end"})


# ---------- inline formatting ----------
def emit(par, text, size=11, base_italic=False):
    for tok in INLINE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r, r.bold = par.add_run(tok[2:-2]), True
        elif tok.startswith("`") and tok.endswith("`"):
            r = par.add_run(tok[1:-1])
            r.font.name, r.font.size = MONO, Pt(size - 1.5)
        elif tok.startswith("*") and tok.endswith("*"):
            r, r.italic = par.add_run(tok[1:-1]), True
        else:
            r = par.add_run(tok)
            r.italic = base_italic
        r.font.name = r.font.name or BODY
        r.font.size = r.font.size or Pt(size)
    return par


# ---------- table building ----------
def add_table(doc, rows):
    ncols = len(rows[0])
    widths = WIDTHS.get(ncols, [6.5 / ncols] * ncols)
    size = 9 if ncols >= 5 else 10
    t = doc.add_table(rows=0, cols=ncols)
    t.style = "Table Grid"
    t.autofit = False
    for ri, cells in enumerate(rows):
        row = t.add_row()
        if ri == 0:
            repeat_header(row)
        for ci, txt in enumerate(cells):
            cell = row.cells[ci]
            par = cell.paragraphs[0]
            par.paragraph_format.space_after = Pt(2)
            if ri == 0:
                run = par.add_run(txt.replace("*", ""))
                run.bold, run.font.name, run.font.size = True, BODY, Pt(size)
                shade(cell._tc, "E8E8E8")
            else:
                emit(par, txt, size=size)
    for row in t.rows:                      # widths must be set per-cell
        for ci, cell in enumerate(row.cells):
            cell.width = Inches(widths[ci])
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


# ---------- document setup ----------
def new_doc():
    doc = Document()
    s = doc.sections[0]
    s.orientation = WD_ORIENT.PORTRAIT
    s.page_width, s.page_height = Inches(8.5), Inches(11)
    s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Inches(1)

    n = doc.styles["Normal"]
    n.font.name, n.font.size = BODY, Pt(11)
    n.paragraph_format.space_after = Pt(6)
    n.paragraph_format.line_spacing = 1.15

    for name, size, caps in (("Heading 1", 13, True),
                             ("Heading 2", 11.5, False),
                             ("Heading 3", 11, False)):
        st = doc.styles[name]
        st.font.name, st.font.size = BODY, Pt(size)
        st.font.bold, st.font.color.rgb = True, RGBColor(0, 0, 0)
        st.font.all_caps = caps
        st.font.italic = (name == "Heading 3")
        st.paragraph_format.space_before = Pt(14 if caps else 10)
        st.paragraph_format.space_after = Pt(5)

    f = s.footer.paragraphs[0]
    f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f.add_run(LEGEND + "   ·   ")
    page_field(f)
    for r in f.runs:
        r.font.name, r.font.size = BODY, Pt(8.5)
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    return doc


# ---------- parser ----------
def convert(md_path, out_path):
    lines = open(md_path, encoding="utf-8").read().split("\n")
    doc = new_doc()
    i, fences = 0, 0

    while i < len(lines):
        line = lines[i]

        if line.startswith("```"):                       # header block / closing notice
            buf = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1
            last = fences > 0
            fences += 1
            for n, b in enumerate(buf):
                p = doc.add_paragraph()
                pf = p.paragraph_format
                pf.space_after = Pt(6 if n == len(buf) - 1 else 0)
                pf.space_before = Pt(0)
                pf.line_spacing = 1.0
                pf.left_indent = pf.right_indent = Inches(0.08)
                r = p.add_run(b)
                r.font.name, r.font.size = MONO, Pt(8 if last else 8.5)
                r.bold = (n == 0 and not last)           # the WORK PRODUCT legend line
                edges = ["left", "right"]
                if n == 0:
                    edges.append("top")
                if n == len(buf) - 1:
                    edges.append("bottom")
                border(p, edges)
            continue

        if line.startswith(">"):                          # best-issue flag
            buf = []
            while i < len(lines) and lines[i].startswith(">"):
                buf.append(re.sub(r"^>\s?", "", lines[i]))
                i += 1
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.left_indent = pf.right_indent = Inches(0.18)
            pf.space_before, pf.space_after = Pt(6), Pt(8)
            emit(p, " ".join(buf).strip())
            shade(p._p, "E8EDF5")
            border(p, ["left"], size="18", color="30508A")
            continue

        if line.startswith("|"):                          # pipe table
            raw = []
            while i < len(lines) and lines[i].startswith("|"):
                raw.append(lines[i])
                i += 1
            rows = [
                [c.strip() for c in r.strip().strip("|").split("|")]
                for r in raw
                if not re.fullmatch(r"\|[\s:|-]+\|", r.strip())
            ]
            add_table(doc, rows)
            continue

        if line.startswith("#### "):
            emit(doc.add_paragraph(style="Heading 3"), line[5:], size=11)
            i += 1
            continue
        if line.startswith("### "):
            emit(doc.add_paragraph(style="Heading 2"), line[4:], size=11.5)
            i += 1
            continue
        if line.startswith("## "):
            emit(doc.add_paragraph(style="Heading 1"), line[3:], size=13)
            i += 1
            continue

        m = ORDERED.match(line)                           # numbered action item
        if m:
            text = line[m.end():]
            i += 1
            while i < len(lines) and re.match(r"^\s{2,}\S", lines[i]):
                text += " " + lines[i].strip()
                i += 1
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(4)
            emit(p, text)
            continue

        if re.match(r"^-\s", line):                       # bullet, possibly wrapped
            text = line[2:]
            i += 1
            while i < len(lines) and re.match(r"^\s{2,}\S", lines[i]):
                text += " " + lines[i].strip()
                i += 1
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(5)
            emit(p, text)
            continue

        if not line.strip():
            i += 1
            continue

        buf = [line]                                      # paragraph, join wrapped lines
        i += 1
        while i < len(lines) and lines[i].strip() and not re.match(r"^(\||>|#|-\s|\d+\.\s|```)", lines[i]):
            buf.append(lines[i])
            i += 1
        emit(doc.add_paragraph(), re.sub(r"\s+", " ", " ".join(buf)).strip())

    doc.save(out_path)
    print("wrote " + out_path)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        sys.exit("usage: analysis_to_docx.py <in.md> <out.docx>")
    convert(sys.argv[1], sys.argv[2])
